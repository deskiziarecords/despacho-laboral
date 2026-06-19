"""
Generador de Demanda Laboral (Word .docx)
==========================================

Genera un documento profesional de Demanda Laboral Mexicana
con todos los datos del expediente, cálculos integrados y
formato apto para impresión y firma.

Autor: Conciliacion Laboral Tijuana - Módulo de Demandas
"""

import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from django.utils import timezone

from .models import Expediente
from .laboral_calculator import calcular_desde_expediente


# ─── Meses en español ──────────────────────────────────────────────────────

MESES_ES = [
    "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


# ─── Plantillas de Demanda (tipos de despido) ──────────────────────────────

PLANTILLAS_INFO = {
    'injustificado': {
        'nombre': 'Despido Injustificado',
        'descripcion': 'El patrón despidió al trabajador sin causa justificada. Reclama indemnización completa (3 meses + antigüedad + prestaciones).',
        'icono': '⚡',
        'recomendado': True,
    },
    'justificado': {
        'nombre': 'Despido Justificado (con responsabilidad al patrón)',
        'descripcion': 'El trabajador rescinde la relación laboral por causas imputables al patrón (falta de pago, maltrato, etc.).',
        'icono': '🛡️',
        'recomendado': False,
    },
    'voluntario': {
        'nombre': 'Renuncia Voluntaria',
        'descripcion': 'El trabajador renunció voluntariamente. Solo reclama prestaciones proporcionales adeudadas (aguinaldo, vacaciones, prima vacacional).',
        'icono': '✍️',
        'recomendado': False,
    },
    'rescision': {
        'nombre': 'Rescisión de la Relación Laboral',
        'descripcion': 'Rescisión imputable al patrón por incumplimiento grave (Art. 51 LFT). Reclama indemnización completa.',
        'icono': '⚖️',
        'recomendado': False,
    },
    'otro': {
        'nombre': 'Otro / Personalizado',
        'descripcion': 'Plantilla genérica para cualquier otra causa de terminación laboral. Edita libremente el contenido.',
        'icono': '📄',
        'recomendado': False,
    },
}


def _fecha_espanol(fecha) -> str:
    """Formatea una fecha en español: '1 de enero de 2024'."""
    if not fecha:
        return "[FECHA]"
    return f"{fecha.day} de {MESES_ES[fecha.month]} de {fecha.year}"


# ─── Estilos ───────────────────────────────────────────────────────────────

TITLE_FONT_SIZE = Pt(14)
SUBTITLE_FONT_SIZE = Pt(12)
SECTION_FONT_SIZE = Pt(11)
BODY_FONT_SIZE = Pt(10)
TABLE_FONT_SIZE = Pt(9.5)

MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)
MARGIN_LEFT = Cm(3)
MARGIN_RIGHT = Cm(2.5)

COLOR_PRIMARY = RGBColor(0x1F, 0x29, 0x37)   # Azul oscuro
COLOR_ACCENT = RGBColor(0x1D, 0x4E, 0xD8)    # Azul acento
COLOR_HEADER_BG = "1F2937"                    # Fondo encabezado tabla
COLOR_ALT_ROW = "F3F4F6"                      # Fila alterna tabla
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_GRAY = RGBColor(0x6B, 0x72, 0x80)


def _configurar_documento(doc: Document) -> None:
    """Configura márgenes, fuente base y orientación."""
    seccion = doc.sections[0]
    seccion.top_margin = MARGIN_TOP
    seccion.bottom_margin = MARGIN_BOTTOM
    seccion.left_margin = MARGIN_LEFT
    seccion.right_margin = MARGIN_RIGHT

    estilo = doc.styles['Normal']
    estilo.font.name = 'Calibri'
    estilo.font.size = BODY_FONT_SIZE
    estilo.font.color.rgb = COLOR_BLACK
    estilo.paragraph_format.space_after = Pt(6)
    estilo.paragraph_format.line_spacing = 1.15


def _agregar_encabezado_tribunal(doc: Document) -> None:
    """Agrega el encabezado con el nombre del tribunal."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRIBUNAL LABORAL COMPETENTE")
    run.bold = True
    run.font.size = TITLE_FONT_SIZE
    run.font.color.rgb = COLOR_PRIMARY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("CIUDAD DE MÉXICO")
    run2.font.size = SUBTITLE_FONT_SIZE
    run2.font.color.rgb = COLOR_GRAY

    # Línea separadora
    p_linea = doc.add_paragraph()
    p_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_linea = p_linea.add_run("─" * 70)
    run_linea.font.color.rgb = COLOR_ACCENT
    run_linea.font.size = Pt(8)
    doc.add_paragraph()


def _celda_sombreada(celda, color: str) -> None:
    """Aplica color de fondo a una celda de tabla."""
    sombreado = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}"/>'
    )
    celda._tc.get_or_add_tcPr().append(sombreado)


def _agregar_materia(doc: Document, expediente: Expediente) -> None:
    """Agrega la materia, tipo de juicio y expediente."""
    tabla = doc.add_table(rows=3, cols=2)
    tabla.style = 'Table Grid'
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    datos = [
        ("MATERIA:", "LABORAL"),
        ("TIPO DE JUICIO:", "ORDINARIO LABORAL"),
        ("N° EXPEDIENTE CONCILIACIÓN:", expediente.folio or "—"),
    ]

    for i, (label, valor) in enumerate(datos):
        celda_label = tabla.cell(i, 0)
        celda_valor = tabla.cell(i, 1)

        p_label = celda_label.paragraphs[0]
        run_label = p_label.add_run(label)
        run_label.bold = True
        run_label.font.size = BODY_FONT_SIZE

        p_valor = celda_valor.paragraphs[0]
        run_valor = p_valor.add_run(valor)
        run_valor.bold = True
        run_valor.font.size = BODY_FONT_SIZE
        run_valor.font.color.rgb = COLOR_ACCENT

        celda_label.width = Cm(5.5)
        celda_valor.width = Cm(7)
        _celda_sombreada(celda_label, COLOR_ALT_ROW)

    doc.add_paragraph()


def _agregar_actor(doc: Document, expediente: Expediente) -> None:
    """Agrega la sección del ACTOR (trabajador)."""
    cliente = expediente.cliente

    p = doc.add_paragraph()
    run = p.add_run("—  A C T O R  —")
    run.bold = True
    run.font.size = SECTION_FONT_SIZE
    run.font.color.rgb = COLOR_PRIMARY

    items = [(cliente.nombre, True)]
    if cliente.direccion_completa:
        items.append((f"Domicilio: {cliente.direccion_completa}", False))
    if cliente.curp:
        items.append((f"CURP: {cliente.curp}", False))
    if cliente.rfc:
        items.append((f"RFC: {cliente.rfc}", False))
    if cliente.telefono:
        items.append((f"Teléfono: {cliente.telefono}", False))

    for texto, negrita in items:
        p = doc.add_paragraph()
        run = p.add_run(f"  {texto}")
        run.bold = negrita
        run.font.size = BODY_FONT_SIZE
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()


def _agregar_demandado(doc: Document, expediente: Expediente) -> None:
    """Agrega la sección del DEMANDADO (patrón/empresa)."""
    cliente = expediente.cliente

    p = doc.add_paragraph()
    run = p.add_run("—  D E M A N D A D O  —")
    run.bold = True
    run.font.size = SECTION_FONT_SIZE
    run.font.color.rgb = COLOR_PRIMARY

    razon_social = cliente.empresa_razon_social or cliente.empresa
    items = [(razon_social or "—", True)]

    partes_dir = []
    if cliente.empresa_calle:
        partes_dir.append(cliente.empresa_calle)
    if cliente.empresa_numero:
        partes_dir.append(f"#{cliente.empresa_numero}")
    if cliente.empresa_colonia:
        partes_dir.append(f"Col. {cliente.empresa_colonia}")
    if cliente.empresa_cp:
        partes_dir.append(f"CP {cliente.empresa_cp}")
    if partes_dir:
        items.append((f"Domicilio: {', '.join(partes_dir)}", False))

    if cliente.empresa_telefono:
        items.append((f"Teléfono: {cliente.empresa_telefono}", False))
    if cliente.empresa_actividad:
        items.append((f"Actividad: {cliente.empresa_actividad}", False))

    for texto, negrita in items:
        p = doc.add_paragraph()
        run = p.add_run(f"  {texto}")
        run.bold = negrita
        run.font.size = BODY_FONT_SIZE
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()


def _narrativa_despido(tipo_despido_key: str) -> str:
    """Devuelve la frase legal correcta según el tipo de despido."""
    narrativas = {
        'injustificado': "el demandado dio por terminada la relación laboral de manera injustificada",
        'justificado': "el actor dio por terminada la relación laboral por causas imputables al demandado",
        'voluntario': "la relación laboral concluyó por renuncia voluntaria del actor",
        'rescision': "el actor se vio en la necesidad de rescindir la relación laboral",
        'otro': "la relación laboral concluyó",
    }
    return narrativas.get(tipo_despido_key,
                          "el demandado dio por terminada la relación laboral de manera injustificada")


def _agregar_hechos(doc: Document, expediente: Expediente, tipo_despido: str = 'injustificado') -> None:
    """Agrega la sección de HECHOS con narrativa legal."""
    cliente = expediente.cliente

    p = doc.add_paragraph()
    run = p.add_run("—  H E C H O S  —")
    run.bold = True
    run.font.size = SECTION_FONT_SIZE
    run.font.color.rgb = COLOR_PRIMARY

    f_ingreso = _fecha_espanol(cliente.fecha_ingreso)
    f_salida = _fecha_espanol(cliente.fecha_salida)
    puesto = cliente.puesto or "[PUESTO DESEMPEÑADO]"
    salario = f"${cliente.salario:,.2f}" if cliente.salario else "[SALARIO]"
    empresa = cliente.empresa_razon_social or cliente.empresa or "[EMPRESA DEMANDADA]"
    folio = expediente.folio or "[FOLIO DE CONCILIACIÓN]"
    f_tramite = _fecha_espanol(expediente.fecha_tramite)
    frase_despido = _narrativa_despido(tipo_despido)

    hechos = f"""
PRIMERO.- El {f_ingreso}, el actor inició su relación laboral con el demandado {empresa}, desempeñando el puesto de {puesto}, con un salario de {salario} mensuales, pagaderos en la forma y términos convenidos.

SEGUNDO.- El {f_salida}, {frase_despido}, violando en perjuicio del actor lo dispuesto por los artículos 46, 47 y 48 de la Ley Federal del Trabajo.

TERCERO.- El actor agotó la instancia conciliatoria ante el Centro de Conciliación Laboral, según consta en el expediente número {folio} de fecha {f_tramite}, sin que se lograra acuerdo conciliatorio alguno, por lo que se expidió la constancia de no conciliación correspondiente.

CUARTO.- A la fecha de presentación de esta demanda, el demandado no ha cubierto al actor el pago de las prestaciones laborales que se reclaman, a pesar de haber sido requerido para ello.
"""

    p_hechos = doc.add_paragraph()
    run_hechos = p_hechos.add_run(hechos.strip())
    run_hechos.font.size = BODY_FONT_SIZE
    doc.add_paragraph()


def _agregar_prestaciones(doc: Document, expediente: Expediente, calculo: dict) -> None:
    """Agrega la sección de PRESTACIONES RECLAMADAS con tabla de montos."""
    p = doc.add_paragraph()
    run = p.add_run("—  P R E S T A C I O N E S   R E C L A M A D A S  —")
    run.bold = True
    run.font.size = SECTION_FONT_SIZE
    run.font.color.rgb = COLOR_PRIMARY

    intro = doc.add_paragraph()
    run_intro = intro.add_run(
        "Con fundamento en lo dispuesto por la Ley Federal del Trabajo, se reclaman "
        "las siguientes prestaciones:"
    )
    run_intro.font.size = BODY_FONT_SIZE
    doc.add_paragraph()

    # Construir filas de la tabla
    rows = [("PRESTACIÓN", "FUNDAMENTO", "IMPORTE")]

    if calculo.get('success'):
        c = calculo
        rows.append(("Aguinaldo Proporcional", "Art. 87 LFT",
                     f"${c['aguinaldo']['monto']:,.2f}"))
        rows.append(("Vacaciones Proporcionales",
                     f"Art. 76 LFT ({c['vacaciones']['dias_segun_antiguedad']} días)",
                     f"${c['vacaciones']['monto']:,.2f}"))
        rows.append(("Prima Vacacional (25%)", "Art. 80 LFT",
                     f"${c['prima_vacacional']['monto']:,.2f}"))
        tope = " (con tope)" if c['prima_antiguedad']['tope_aplicado'] else ""
        rows.append(("Prima de Antigüedad", f"Art. 162 LFT{tope}",
                     f"${c['prima_antiguedad']['monto']:,.2f}"))
        rows.append(("Indemnización Constitucional (3 meses)", "Art. 50 LFT",
                     f"${c['indemnizacion']['monto']:,.2f}"))
        rows.append(("", "TOTAL:", f"${c['total']:,.2f}"))
    else:
        rows.append(("Aguinaldo Proporcional", "Art. 87 LFT", "—"))
        rows.append(("Vacaciones Proporcionales", "Art. 76 LFT", "—"))
        rows.append(("Prima Vacacional (25%)", "Art. 80 LFT", "—"))
        rows.append(("Prima de Antigüedad", "Art. 162 LFT", "—"))
        rows.append(("Indemnización Constitucional (3 meses)", "Art. 50 LFT", "—"))
        if expediente.monto_reclamado:
            rows.append(("", "MONTO RECLAMADO:", f"${expediente.monto_reclamado:,.2f}"))
        else:
            rows.append(("", "MONTO RECLAMADO:", "—"))

    tabla = doc.add_table(rows=len(rows), cols=3)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = 'Table Grid'

    for i, (col1, col2, col3) in enumerate(rows):
        row = tabla.rows[i]
        es_encabezado = (i == 0)
        es_total = (i == len(rows) - 1)

        for j, celda in enumerate(row.cells):
            p_celda = celda.paragraphs[0]
            texto = [col1, col2, col3][j]

            # Alineación
            if j == 2 or (j == 1 and es_total):
                p_celda.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif j == 0:
                p_celda.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p_celda.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run_celda = p_celda.add_run(texto)
            run_celda.font.size = TABLE_FONT_SIZE
            run_celda.bold = es_encabezado or es_total

            if es_encabezado:
                run_celda.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _celda_sombreada(celda, COLOR_HEADER_BG)
            elif es_total:
                run_celda.font.color.rgb = COLOR_ACCENT
            elif i % 2 == 0:
                _celda_sombreada(celda, COLOR_ALT_ROW)

    doc.add_paragraph()


def _agregar_derecho(doc: Document, tipo_despido: str = 'injustificado') -> None:
    """Agrega la sección de FUNDAMENTOS DE DERECHO adaptados al tipo de despido."""
    p = doc.add_paragraph()
    run = p.add_run("—  F U N D A M E N T O S   D E   D E R E C H O  —")
    run.bold = True
    run.font.size = SECTION_FONT_SIZE
    run.font.color.rgb = COLOR_PRIMARY

    articulos = _fundamentos_derecho_docx(tipo_despido)

    for art in articulos:
        p = doc.add_paragraph()
        run = p.add_run(f"  • {art}")
        run.font.size = BODY_FONT_SIZE
        p.paragraph_format.space_after = Pt(1)

    doc.add_paragraph()


def _agregar_puntos_petitorios(doc: Document, expediente: Expediente, calculo: dict) -> None:
    """Agrega los PUNTOS PETITORIOS."""
    p = doc.add_paragraph()
    run = p.add_run("—  P U N T O S   P E T I T O R I O S  —")
    run.bold = True
    run.font.size = SECTION_FONT_SIZE
    run.font.color.rgb = COLOR_PRIMARY

    if calculo.get('success') and calculo['total'] > 0:
        total_str = f"${calculo['total']:,.2f}"
    elif expediente.monto_reclamado:
        total_str = f"${expediente.monto_reclamado:,.2f}"
    else:
        total_str = "la cantidad que resulte"

    petitorios = [
        "PRIMERO.- Se declare que existió una relación laboral entre el actor y el demandado.",
        f"SEGUNDO.- Se condene al demandado al pago de {total_str} por concepto de las prestaciones laborales detalladas en el cuerpo de esta demanda.",
        "TERCERO.- Se ordene el pago de los salarios caídos que se sigan generando hasta la fecha en que se cumpla la sentencia.",
        "CUARTO.- Se condene al demandado al pago de los gastos y costas que se originen con motivo del presente juicio.",
    ]

    for pet in petitorios:
        p = doc.add_paragraph()
        run = p.add_run(f"  {pet}")
        run.font.size = BODY_FONT_SIZE
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()


def _agregar_firma(doc: Document, expediente: Expediente) -> None:
    """Agrega la sección de FIRMA."""
    p = doc.add_paragraph()
    run = p.add_run("—  F I R M A  —")
    run.bold = True
    run.font.size = SECTION_FONT_SIZE
    run.font.color.rgb = COLOR_PRIMARY

    # Fecha en español
    hoy = timezone.now()
    fecha_str = _fecha_espanol(hoy)

    p = doc.add_paragraph()
    run = p.add_run(f"Presentado en la Ciudad de México, a los {fecha_str}.")
    run.font.size = BODY_FONT_SIZE

    doc.add_paragraph()

    # Línea de firma
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 40)
    run.font.size = BODY_FONT_SIZE
    run.font.color.rgb = COLOR_GRAY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(expediente.cliente.nombre)
    run2.bold = True
    run2.font.size = SUBTITLE_FONT_SIZE

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Actor")
    run3.font.size = BODY_FONT_SIZE
    run3.font.color.rgb = COLOR_GRAY

    doc.add_paragraph()

    # Asesor
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    asesor = expediente.asesor.get_full_name() or expediente.asesor.username
    run4 = p4.add_run(f"Asesor jurídico: {asesor}")
    run4.font.size = Pt(8)
    run4.font.color.rgb = COLOR_GRAY


def _agregar_pie_generacion(doc: Document, expediente: Expediente) -> None:
    """Agrega metadata de generación al final del documento."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 70)
    run.font.size = Pt(6)
    run.font.color.rgb = COLOR_GRAY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ahora = timezone.now().strftime('%d/%m/%Y %H:%M')
    asesor = expediente.asesor.get_full_name() or expediente.asesor.username
    run2 = p2.add_run(
        f"Documento generado el {ahora} por {asesor} | "
        f"Exp: {expediente.numero} | "
        f"Sistema de Gestión Laboral"
    )
    run2.font.size = Pt(7)
    run2.font.color.rgb = COLOR_GRAY


def generar_demanda_word(expediente: Expediente, desde_cero=True,
                          tipo_despido_override: str | None = None) -> Document:
    """
    Genera un documento Word de Demanda Laboral completo.

    Args:
        expediente: Instancia de Expediente con datos del cliente
        desde_cero: Si es True, genera el documento completo desde cero.
                    Si es False, solo configura márgenes y fuente (para usar con html_a_docx).
        tipo_despido_override: Si se proporciona, usa este tipo de despido
                               en lugar del que tiene el expediente.

    Returns:
        Documento python-docx listo para guardar/enviar
    """
    doc = Document()
    _configurar_documento(doc)

    if desde_cero:
        tipo_despido = tipo_despido_override or expediente.tipo_despido or 'injustificado'
        calculo = calcular_desde_expediente(expediente)
        _agregar_encabezado_tribunal(doc)
        _agregar_materia(doc, expediente)
        _agregar_actor(doc, expediente)
        _agregar_demandado(doc, expediente)
        _agregar_hechos(doc, expediente, tipo_despido)
        _agregar_prestaciones(doc, expediente, calculo)
        _agregar_derecho(doc, tipo_despido)
        _agregar_puntos_petitorios(doc, expediente, calculo)
        _agregar_firma(doc, expediente)
        _agregar_pie_generacion(doc, expediente)

    return doc


# ════════════════════════════════════════════════════════════════════════
# GENERADOR HTML (para el editor WYSIWYG)
# ════════════════════════════════════════════════════════════════════════

def _fecha_espanol_html(fecha) -> str:
    """Igual que _fecha_espanol pero para HTML."""
    if not fecha:
        return "[FECHA]"
    return f"{fecha.day} de {MESES_ES[fecha.month]} de {fecha.year}"


def _narrativa_despido_html(tipo_despido_key: str) -> str:
    """Igual que _narrativa_despido pero para HTML."""
    narrativas = {
        'injustificado': "el demandado dio por terminada la relación laboral de manera injustificada",
        'justificado': "el actor dio por terminada la relación laboral por causas imputables al demandado",
        'voluntario': "la relación laboral concluyó por renuncia voluntaria del actor",
        'rescision': "el actor se vio en la necesidad de rescindir la relación laboral",
        'otro': "la relación laboral concluyó",
    }
    return narrativas.get(tipo_despido_key,
                          "el demandado dio por terminada la relación laboral de manera injustificada")


def _fundamentos_derecho_html(tipo_despido_key: str) -> str:
    """Genera los fundamentos de derecho adaptados al tipo de despido."""
    articulos_base = [
        "Artículo 84 LFT — Salario integrado",
        "Artículo 87 LFT — Aguinaldo anual (15 días mínimo)",
        "Artículo 76 LFT — Vacaciones",
        "Artículo 79 LFT — Prima vacacional (mínimo 25%)",
        "Artículo 80 LFT — Pago de prima vacacional",
        "Artículo 518 LFT — Procedimiento ordinario laboral",
    ]

    if tipo_despido_key in ('injustificado', 'rescision'):
        adicionales = [
            "Artículo 46 LFT — Terminación de la relación laboral",
            "Artículo 47 LFT — Causas de rescisión sin responsabilidad",
            "Artículo 48 LFT — Indemnización por despido injustificado",
            "Artículo 49 LFT — Exención de responsabilidad al patrón",
            "Artículo 50 LFT — Indemnización de 3 meses",
            "Artículo 162 LFT — Prima de antigüedad",
        ]
    elif tipo_despido_key == 'justificado':
        adicionales = [
            "Artículo 46 LFT — Terminación de la relación laboral",
            "Artículo 51 LFT — Causas de rescisión imputables al patrón",
            "Artículo 52 LFT — Aviso de rescisión al patrón",
            "Artículo 48 LFT — Indemnización por despido injustificado",
            "Artículo 50 LFT — Indemnización de 3 meses",
            "Artículo 162 LFT — Prima de antigüedad",
        ]
    elif tipo_despido_key == 'voluntario':
        adicionales = [
            "Artículo 53 LFT — Causas de terminación de la relación laboral (Fracción I)",
        ]
    else:
        adicionales = [
            "Artículo 46 LFT — Terminación de la relación laboral",
        ]

    return "\n".join(f"<p>• {art}</p>" for art in articulos_base + adicionales)


def _fundamentos_derecho_docx(tipo_despido_key: str) -> list:
    """Genera lista de fundamentos de derecho adaptados al tipo de despido para docx."""
    articulos_base = [
        "Artículo 84 LFT — Salario integrado",
        "Artículo 87 LFT — Aguinaldo anual (15 días mínimo)",
        "Artículo 76 LFT — Vacaciones",
        "Artículo 79 LFT — Prima vacacional (mínimo 25%)",
        "Artículo 80 LFT — Pago de prima vacacional",
        "Artículo 518 LFT — Procedimiento ordinario laboral",
    ]

    if tipo_despido_key in ('injustificado', 'rescision'):
        adicionales = [
            "Artículo 46 LFT — Terminación de la relación laboral",
            "Artículo 47 LFT — Causas de rescisión sin responsabilidad",
            "Artículo 48 LFT — Indemnización por despido injustificado",
            "Artículo 49 LFT — Exención de responsabilidad al patrón",
            "Artículo 50 LFT — Indemnización de 3 meses",
            "Artículo 162 LFT — Prima de antigüedad",
        ]
    elif tipo_despido_key == 'justificado':
        adicionales = [
            "Artículo 46 LFT — Terminación de la relación laboral",
            "Artículo 51 LFT — Causas de rescisión imputables al patrón",
            "Artículo 52 LFT — Aviso de rescisión al patrón",
            "Artículo 48 LFT — Indemnización por despido injustificado",
            "Artículo 50 LFT — Indemnización de 3 meses",
            "Artículo 162 LFT — Prima de antigüedad",
        ]
    elif tipo_despido_key == 'voluntario':
        adicionales = [
            "Artículo 53 LFT — Causas de terminación de la relación laboral (Fracción I)",
        ]
    else:
        adicionales = [
            "Artículo 46 LFT — Terminación de la relación laboral",
        ]

    return articulos_base + adicionales


def generar_demanda_html(expediente: Expediente, tipo_despido_override: str | None = None) -> str:
    """
    Genera el contenido de la Demanda Laboral como HTML
    para ser usado en el editor WYSIWYG (Quill.js).

    Args:
        expediente: Instancia de Expediente con datos del cliente
        tipo_despido_override: Si se proporciona, usa este tipo de despido
                               en lugar del que tiene el expediente.
    """
    cliente = expediente.cliente
    calculo = calcular_desde_expediente(expediente)
    hoy = timezone.now()
    asesor = expediente.asesor.get_full_name() or expediente.asesor.username
    ahora_str = hoy.strftime('%d/%m/%Y %H:%M')

    tipo_despido = tipo_despido_override or expediente.tipo_despido or 'injustificado'

    # ─── Fechas ───
    f_ingreso = _fecha_espanol_html(cliente.fecha_ingreso)
    f_salida = _fecha_espanol_html(cliente.fecha_salida)
    puesto = cliente.puesto or "[PUESTO DESEMPEÑADO]"
    salario = f"${cliente.salario:,.2f}" if cliente.salario else "[SALARIO]"
    empresa = cliente.empresa_razon_social or cliente.empresa or "[EMPRESA DEMANDADA]"
    folio = expediente.folio or "[FOLIO DE CONCILIACIÓN]"
    f_tramite = _fecha_espanol_html(expediente.fecha_tramite)
    frase_despido = _narrativa_despido_html(tipo_despido)
    fecha_str = _fecha_espanol_html(hoy)

    # ─── Prestaciones ───
    prestaciones_rows = ""
    if calculo.get('success'):
        c = calculo
        prestaciones_rows = f"""
        <tr><td>Aguinaldo Proporcional</td><td>Art. 87 LFT</td><td style="text-align:right">${c['aguinaldo']['monto']:,.2f}</td></tr>
        <tr><td>Vacaciones Proporcionales</td><td>Art. 76 LFT ({c['vacaciones']['dias_segun_antiguedad']} días)</td><td style="text-align:right">${c['vacaciones']['monto']:,.2f}</td></tr>
        <tr><td>Prima Vacacional (25%)</td><td>Art. 80 LFT</td><td style="text-align:right">${c['prima_vacacional']['monto']:,.2f}</td></tr>
        <tr><td>Prima de Antigüedad</td><td>Art. 162 LFT{' (con tope)' if c['prima_antiguedad']['tope_aplicado'] else ''}</td><td style="text-align:right">${c['prima_antiguedad']['monto']:,.2f}</td></tr>
        <tr><td>Indemnización Constitucional (3 meses)</td><td>Art. 50 LFT</td><td style="text-align:right">${c['indemnizacion']['monto']:,.2f}</td></tr>
        <tr style="font-weight:bold;border-top:2px solid #000"><td></td><td style="text-align:right">TOTAL:</td><td style="text-align:right">${c['total']:,.2f}</td></tr>
        """
    else:
        total_str = f"${expediente.monto_reclamado:,.2f}" if expediente.monto_reclamado else "—"
        prestaciones_rows = f"""
        <tr><td>Aguinaldo Proporcional</td><td>Art. 87 LFT</td><td style="text-align:right">—</td></tr>
        <tr><td>Vacaciones Proporcionales</td><td>Art. 76 LFT</td><td style="text-align:right">—</td></tr>
        <tr><td>Prima Vacacional (25%)</td><td>Art. 80 LFT</td><td style="text-align:right">—</td></tr>
        <tr><td>Prima de Antigüedad</td><td>Art. 162 LFT</td><td style="text-align:right">—</td></tr>
        <tr><td>Indemnización Constitucional (3 meses)</td><td>Art. 50 LFT</td><td style="text-align:right">—</td></tr>
        <tr style="font-weight:bold;border-top:2px solid #000"><td></td><td style="text-align:right">MONTO RECLAMADO:</td><td style="text-align:right">{total_str}</td></tr>
        """

    total_petitorio = f"${calculo['total']:,.2f}" if calculo.get('success') and calculo['total'] > 0 else (f"${expediente.monto_reclamado:,.2f}" if expediente.monto_reclamado else "la cantidad que resulte")

    # ─── Datos actor ───
    actor_direccion = cliente.direccion_completa
    actor_items = [f"<strong>{cliente.nombre}</strong>"]
    if actor_direccion:
        actor_items.append(f"Domicilio: {actor_direccion}")
    if cliente.curp:
        actor_items.append(f"CURP: {cliente.curp}")
    if cliente.rfc:
        actor_items.append(f"RFC: {cliente.rfc}")
    if cliente.telefono:
        actor_items.append(f"Teléfono: {cliente.telefono}")

    # ─── Datos demandado ───
    razon_social = cliente.empresa_razon_social or cliente.empresa
    demandado_items = [f"<strong>{razon_social or '—'}</strong>"]
    partes_dir = []
    if cliente.empresa_calle:
        partes_dir.append(cliente.empresa_calle)
    if cliente.empresa_numero:
        partes_dir.append(f"#{cliente.empresa_numero}")
    if cliente.empresa_colonia:
        partes_dir.append(f"Col. {cliente.empresa_colonia}")
    if cliente.empresa_cp:
        partes_dir.append(f"CP {cliente.empresa_cp}")
    if partes_dir:
        demandado_items.append(f"Domicilio: {', '.join(partes_dir)}")
    if cliente.empresa_telefono:
        demandado_items.append(f"Teléfono: {cliente.empresa_telefono}")
    if cliente.empresa_actividad:
        demandado_items.append(f"Actividad: {cliente.empresa_actividad}")

    html = f"""
<h2 style="text-align:center;color:#1F2937;">TRIBUNAL LABORAL COMPETENTE</h2>
<p style="text-align:center;color:#6B7280;">CIUDAD DE MÉXICO</p>
<hr style="border:none;border-top:1px solid #1D4ED8;width:70%;margin:10px auto;">

<table style="width:100%;border-collapse:collapse;margin:15px 0;">
    <tr><td style="background:#F3F4F6;padding:5px 8px;border:1px solid #ccc;font-weight:bold;">MATERIA:</td><td style="padding:5px 8px;border:1px solid #ccc;font-weight:bold;color:#1D4ED8;">LABORAL</td></tr>
    <tr><td style="background:#F3F4F6;padding:5px 8px;border:1px solid #ccc;font-weight:bold;">TIPO DE JUICIO:</td><td style="padding:5px 8px;border:1px solid #ccc;font-weight:bold;color:#1D4ED8;">ORDINARIO LABORAL</td></tr>
    <tr><td style="background:#F3F4F6;padding:5px 8px;border:1px solid #ccc;font-weight:bold;">N° EXPEDIENTE CONCILIACIÓN:</td><td style="padding:5px 8px;border:1px solid #ccc;font-weight:bold;color:#1D4ED8;">{expediente.folio or '—'}</td></tr>
</table>

<h3 style="color:#1F2937;">—  A C T O R  —</h3>
"""

    for item in actor_items:
        html += f"<p style='margin:2px 0;'>{item}</p>\n"

    html += f"""

<h3 style="color:#1F2937;">—  D E M A N D A D O  —</h3>
"""

    for item in demandado_items:
        html += f"<p style='margin:2px 0;'>{item}</p>\n"

    html += f"""

<h3 style="color:#1F2937;">—  H E C H O S  —</h3>

<p><strong>PRIMERO.-</strong> El {f_ingreso}, el actor inició su relación laboral con el demandado {empresa}, desempeñando el puesto de {puesto}, con un salario de {salario} mensuales, pagaderos en la forma y términos convenidos.</p>

<p><strong>SEGUNDO.-</strong> El {f_salida}, {frase_despido}, violando en perjuicio del actor lo dispuesto por los artículos 46, 47 y 48 de la Ley Federal del Trabajo.</p>

<p><strong>TERCERO.-</strong> El actor agotó la instancia conciliatoria ante el Centro de Conciliación Laboral, según consta en el expediente número {folio} de fecha {f_tramite}, sin que se lograra acuerdo conciliatorio alguno, por lo que se expidió la constancia de no conciliación correspondiente.</p>

<p><strong>CUARTO.-</strong> A la fecha de presentación de esta demanda, el demandado no ha cubierto al actor el pago de las prestaciones laborales que se reclaman, a pesar de haber sido requerido para ello.</p>

<h3 style="color:#1F2937;">—  P R E S T A C I O N E S   R E C L A M A D A S  —</h3>

<p>Con fundamento en lo dispuesto por la Ley Federal del Trabajo, se reclaman las siguientes prestaciones:</p>

<table style="width:100%;border-collapse:collapse;margin:15px 0;">
    <thead>
        <tr style="background:#1F2937;color:white;">
            <th style="padding:6px 8px;border:1px solid #ccc;text-align:left;">PRESTACIÓN</th>
            <th style="padding:6px 8px;border:1px solid #ccc;text-align:center;">FUNDAMENTO</th>
            <th style="padding:6px 8px;border:1px solid #ccc;text-align:right;">IMPORTE</th>
        </tr>
    </thead>
    <tbody>
{prestaciones_rows}
    </tbody>
</table>

<h3 style="color:#1F2937;">—  F U N D A M E N T O S   D E   D E R E C H O  —</h3>

{_fundamentos_derecho_html(tipo_despido)}

<h3 style="color:#1F2937;">—  P U N T O S   P E T I T O R I O S  —</h3>

<p><strong>PRIMERO.-</strong> Se declare que existió una relación laboral entre el actor y el demandado.</p>
<p><strong>SEGUNDO.-</strong> Se condene al demandado al pago de {total_petitorio} por concepto de las prestaciones laborales detalladas en el cuerpo de esta demanda.</p>
<p><strong>TERCERO.-</strong> Se ordene el pago de los salarios caídos que se sigan generando hasta la fecha en que se cumpla la sentencia.</p>
<p><strong>CUARTO.-</strong> Se condene al demandado al pago de los gastos y costas que se originen con motivo del presente juicio.</p>

<h3 style="color:#1F2937;">—  F I R M A  —</h3>

<p>Presentado en la Ciudad de México, a los {fecha_str}.</p>

<br>
<p style="text-align:center;">________________________________________</p>
<p style="text-align:center;font-weight:bold;font-size:14px;">{cliente.nombre}</p>
<p style="text-align:center;color:#6B7280;">Actor</p>

<br>
<p style="text-align:right;font-size:10px;color:#6B7280;">Asesor jurídico: {asesor}</p>

<hr style="border:none;border-top:1px solid #ccc;width:70%;margin:10px auto;">
<p style="text-align:center;font-size:9px;color:#6B7280;">Documento generado el {ahora_str} por {asesor} | Exp: {expediente.numero} | Sistema de Gestión Laboral</p>
"""

    return html


def html_a_docx(html: str, doc: Document) -> None:
    """
    Convierte HTML básico (generado por Quill.js) a elementos de python-docx.
    Se agrega al documento existente.

    Soporta: p, h1-h3, strong, em, u, br, ol/ul/li, table.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Limpiar el HTML de Quill: reemplazar br
    html = html.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')

    def _estilo_parrafo(paragraph, align_text=None):
        """Aplica formato básico a un párrafo."""
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.15
        if align_text == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align_text == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _agregar_run(paragraph, texto, bold=False, italic=False, underline=False, size=None, color=None):
        """Agrega un run con formato al párrafo."""
        run = paragraph.add_run(texto)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return run

    def _procesar_inline(text, paragraph):
        """
        Procesa etiquetas inline (strong, em, u, s) dentro de un párrafo.
        Maneja múltiples etiquetas en un mismo párrafo.
        """
        remaining = text
        while remaining:
            # Buscar la siguiente etiqueta de apertura
            tag_match = re.search(r'<(strong|b|em|i|u|s)>(.*?)</\1>', remaining, re.DOTALL)
            if tag_match:
                # Texto antes de la etiqueta
                before = remaining[:tag_match.start()]
                if before:
                    paragraph.add_run(before)
                # La etiqueta con formato
                tag = tag_match.group(1)
                content = tag_match.group(2)
                run = paragraph.add_run(content)
                if tag in ('strong', 'b'):
                    run.bold = True
                if tag in ('em', 'i'):
                    run.italic = True
                if tag == 'u':
                    run.underline = True
                if tag == 's':
                    run.font.strike = True
                # Avanzar
                remaining = remaining[tag_match.end():]
            else:
                # No hay más etiquetas, agregar el resto
                if remaining:
                    paragraph.add_run(remaining)
                break

    # Procesar líneas del HTML
    lines = html.split('\n')
    i = 0
    in_list = False
    list_type = None
    current_list_items = []
    in_table = False
    table_data = []

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Tablas
        if '<table' in line:
            in_table = True
            table_data = []
            table_row = []
            # Buscar celdas dentro de la tabla
            table_html = ''
            while i < len(lines) and '</table>' not in lines[i]:
                table_html += lines[i] + '\n'
                i += 1
            table_html += lines[i] + '\n'  # </table>
            i += 1

            # Extraer filas
            rows = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL | re.IGNORECASE)
            if rows:
                parsed_rows = []
                for row_html in rows:
                    # Encontrar cada celda individualmente para saber si es <th> o <td>
                    cell_pattern = re.findall(r'<(th|td)[^>]*>(.*?)</\1>', row_html, re.DOTALL | re.IGNORECASE)
                    parsed_cells = []
                    for tag, content in cell_pattern:
                        # Limpiar etiquetas HTML internas
                        clean_text = re.sub(r'<[^>]+>', '', content).strip()
                        parsed_cells.append({'text': clean_text, 'is_header': tag == 'th'})
                    if parsed_cells:
                        parsed_rows.append(parsed_cells)

                if parsed_rows:
                    num_cols = max(len(r) for r in parsed_rows)
                    table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
                    table.style = 'Table Grid'
                    for ri, row_data in enumerate(parsed_rows):
                        for ci, cell_data in enumerate(row_data):
                            if ci < num_cols:
                                cell = table.rows[ri].cells[ci]
                                cell.text = ''
                                p = cell.paragraphs[0]
                                run = p.add_run(cell_data['text'])
                                run.font.size = Pt(9.5)
                                if cell_data['is_header']:
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                    from docx.oxml.ns import nsdecls
                                    from docx.oxml import parse_xml
                                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F2937"/>')
                                    cell._tc.get_or_add_tcPr().append(shading)
                                # Alinear a la derecha la última columna
                                if ci == num_cols - 1:
                                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_paragraph()  # espacio después de tabla
            continue

        # Listas
        if '<ol>' in line or '<ul>' in line:
            in_list = True
            list_type = 'ol' if '<ol>' in line else 'ul'
            i += 1
            continue
        if '</ol>' in line or '</ul>' in line:
            # Escribir items de lista
            for idx, item_text in enumerate(current_list_items):
                p = doc.add_paragraph()
                _estilo_parrafo(p)
                prefix = f"{idx + 1}. " if list_type == 'ol' else "  • "
                _procesar_inline(f"{prefix}{item_text}", p)
            current_list_items = []
            in_list = False
            list_type = None
            i += 1
            continue
        if in_list:
            li_match = re.search(r'<li>(.*?)</li>', line)
            if li_match:
                current_list_items.append(li_match.group(1))
            i += 1
            continue

        # Encabezados
        h_match = re.match(r'<(h[123])[^>]*>(.*?)</\1>', line, re.DOTALL)
        if h_match:
            tag = h_match.group(1)
            content = h_match.group(2)
            clean_content = re.sub(r'<[^>]+>', '', content).strip()
            p = doc.add_paragraph()
            size_map = {'h1': 14, 'h2': 12, 'h3': 11}
            _agregar_run(p, clean_content, bold=True, size=size_map.get(tag, 11), color=(0x1F, 0x29, 0x37))
            _estilo_parrafo(p)
            i += 1
            continue

        # Párrafos normales
        p_match = re.match(r'<p[^>]*>(.*?)</p>', line, re.DOTALL)
        if p_match:
            p_content = p_match.group(1).strip()
            # Detectar align
            align = None
            align_match = re.search(r'style="[^"]*text-align:\s*(center|right)"', line)
            if align_match:
                align = align_match.group(1)

            p = doc.add_paragraph()
            _estilo_parrafo(p, align)
            _procesar_inline(p_content, p)
            i += 1
            continue

        # Línea horizontal
        if '<hr' in line:
            p = doc.add_paragraph()
            _agregar_run(p, "─" * 70, size=8, color=(0x6B, 0x72, 0x80))
            i += 1
            continue

        # Cualquier otra línea
        clean = re.sub(r'<[^>]+>', '', line).strip()
        if clean:
            p = doc.add_paragraph()
            _estilo_parrafo(p)
            p.add_run(clean)
        i += 1
